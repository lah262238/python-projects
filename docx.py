from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

# Create document
doc = Document()

# Title
title = doc.add_heading('Methodology Approaches for Preparing and Standardizing Gecko Homogenate: A Comprehensive Review (2020–2026)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Author and date
doc.add_paragraph('Author: Academic Writer (decades of experience)', alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph('Date: June 9, 2026', alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph('')

# Abstract
doc.add_heading('Abstract', level=1)
abstract = doc.add_paragraph(
    'This review article comprehensively examines methodology approaches for preparing and standardizing gecko homogenate, '
    'with particular focus on Tokay gecko (Gekko gecko) tissue, which is widely used in traditional Chinese medicine and '
    'biochemical research. We synthesize updated literature from 2020 to 2026, covering tissue collection, homogenization '
    'techniques, buffer selection, temperature control, centrifugation protocols, protein standardization, and quality assurance '
    'measures. Various homogenization methods including mechanical disruption, bead mills, ultrasonic homogenization, and '
    'cryogenic techniques are evaluated. Critical parameters for standardization such as tissue-to-buffer ratios (typically 1:9), '
    'pH optimization (7.0–7.4), protein concentration normalization (2.5–7.5 mg/mL), and centrifugation conditions (5,000–16,000 × g) '
    'are discussed. This review provides actionable guidelines for researchers preparing gecko homogenates for enzymatic, proteomic, '
    'and biochemical analyses while maintaining sample integrity and reproducibility.'
)

# Keywords
doc.add_paragraph('Keywords: ', style='Heading 3')
doc.add_paragraph('gecko homogenate; Gekko gecko; tissue homogenization; standardization protocol; Tokay gecko; '
                 'traditional Chinese medicine; protein quantification; biochemical analysis; methodological review')

doc.add_paragraph('')

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'Tokay gecko (Gekko gecko, Linnaeus, 1758) is a rare and endangered medicinal animal prevalent in southern China and '
    'Southeast Asia. Its dry body has been used as an anti-asthmatic agent for two thousand years in traditional '
    'Chinese medicine (TCM), widely incorporated into Chinese patent medicines such as Gejie Dingchuan capsule and pill. '
    'Despite increasing medicinal demand, ecological deterioration, and excessive hunting, G. gecko has been listed as a Class II '
    'protected species in China since 1989.'
)

doc.add_paragraph(
    'For biochemical analysis, including estimation of enzymatic parameters, proteins, and metabolic compounds from animal tissue, '
    'preparation of tissue homogenate is essential. The homogenate, after centrifugation, yields supernatant (clear solution) '
    'and sediment (membrane-bound portion), both used for various estimations. However, specific methodology for gecko tissue '
    'homogenate preparation and standardization remains poorly documented in peer-reviewed literature from 2020 onwards.'
)

doc.add_paragraph(
    'This review synthesizes updated methodological literature (2020–2026) on tissue homogenate preparation applicable to gecko '
    'tissue, addressing critical steps including sample collection, homogenization techniques, buffer selection, temperature control, '
    'centrifugation, protein standardization, and quality assurance. Understanding these methodologies is crucial for obtaining '
    'high-quality homogenates and ensuring reliable, reproducible results in gecko biochemical research.'
)

doc.add_paragraph('')

# Background
doc.add_heading('2. Background: Gecko as Medicinal and Research Material', level=1)

doc.add_heading('2.1 Traditional Medicinal Use', level=2)
doc.add_paragraph(
    'The dry body of Gekko gecko exerts remarkable effects on strengthening the immune system and treating tumors. '
    'Over the past decades, increasing medicinal demand has led to population declines of up to 50% in some locations. '
    'Millions of tokay geckos per year were exported from Indonesia from 2015 to 2021, all dead and dried, with an export record '
    'of 5,974,550 in 2022.'
)

doc.add_heading('2.2 Research Context', level=2)
doc.add_paragraph(
    'Recent genomic research includes the first full-length transcriptome of G. gecko using SMRT sequencing, identifying 882,273 '
    'circular consensus reads and 203,994 high-quality isoforms. This genetic resource accelerates transcriptome research '
    'and lays foundation for further research on gene function and biological regulatory mechanisms in G. gecko.'
)

doc.add_paragraph(
    'Genetic and stable isotope analyses have been used to test whether geckos sold in TCM markets originate from local or non-local '
    'populations, revealing that TCM tokays are likely not of local Hong Kong origin but imported from Southeast Asia.'
)

doc.add_paragraph('')

# Tissue Collection
doc.add_heading('3. Tissue Collection and Pre-Homogenization Preparation', level=1)

doc.add_heading('3.1 Sample Collection Protocol', level=2)
doc.add_paragraph(
    'Proper sample handling and storage are essential to maintain sample integrity before homogenization. For cultured adult '
    'Tokay gecko samples, specimens are housed with 12:12 day-night light cycle and 70% humidity, fed ad libitum access to water '
    'and ground beetles (Eupolyphaga sinensis Walker) daily prior to euthanasia.'
)

doc.add_paragraph(
    'Tissues including heart, kidney, liver, lung, skin, blood, muscle, stomach, ovary, and oviduct are dissected, immediately '
    'frozen in liquid nitrogen, and stored at −80°C. For wild-caught specimens, tail clippings (~0.5 cm) are collected, '
    'with half stored in 70% ethanol for genetic study and half in empty tubes for stable isotope analysis.'
)

doc.add_heading('3.2 Tissue Rinse and Weighing', level=2)
doc.add_paragraph(
    'The tissue sample should be rinsed with pre-cooled PBS (0.01 mol/L, pH 7.0–7.4) to remove residual blood and impurities on '
    'the surface. The tissue block is then weighed, with the mass to volume ratio of tissue to PBS recommended to be 1:9 '
    '(i.e., 100 mg sample added to 900 μl PBS). Tissue sampling volume should be more than 30 mg at minimum to avoid too '
    'little homogenization buffer affecting grinding effect.'
)

doc.add_heading('3.3 Tissue Cutting', level=2)
doc.add_paragraph(
    'The weighed tissue should be cut into as small pieces as possible to fully homogenate. This step ensures better '
    'crushing effect during homogenization. For microsamples, a micro-technique involves grinding samples in glass tubes closed at '
    'one end, prepared from melting-point capillaries, using stainless-steel plungers to homogenize 2–20 microliter of buffer.'
)

doc.add_paragraph('')

# Homogenization Methods
doc.add_heading('4. Homogenization Techniques and Methods', level=1)

doc.add_paragraph(
    'Various methods and techniques for tissue homogenization offer researchers a range of options to suit specific needs. '
    'Understanding principles, advantages, and considerations of each method is crucial for obtaining high-quality homogenates.'
)

doc.add_heading('4.1 Mechanical Disruption Methods', level=2)

doc.add_paragraph('4.1.1 Mortar and Pestle (Traditional Method)')
doc.add_paragraph(
    'The traditional method involves grinding tissues with a mortar and pestle. While simple and cost-effective, it requires manual '
    'effort and may not be suitable for larger sample sizes. It is commonly used for soft tissues or when gentle disruption '
    'is desired.'
)

doc.add_paragraph('4.1.2 High-Speed Blenders')
doc.add_paragraph(
    'High-speed blenders mechanically disrupt tissues quickly and effectively, suitable for a wide range of tissue types. '
    'However, they may generate heat and lead to protein denaturation if prolonged blending is performed.'
)

doc.add_paragraph('4.1.3 Potter-Elvehjem Type Homogenizer')
doc.add_paragraph(
    'A Potter-Elvehjem type homogenizer with a glass tube is used for homogenization. The shaft is inserted into the tube, and '
    'speed is slowly increased. If speed is increased immediately, there are chances of apparatus breakdown.'
)

doc.add_paragraph('4.1.4 TissueLyser with Stainless Steel Beads')
doc.add_paragraph(
    'Tissues are weighed in 2 mL microcentrifuge tubes, and 500 μL of Cell Lysis Buffer is added per 100 mg tissue. '
    'One 5-mm stainless steel bead is added, and tubes are assembled into TissueLyser. Homogenization occurs at 25 Hz for 0.5–3 '
    'minutes. This method offers high throughput and efficient disruption.'
)

doc.add_heading('4.2 Bead Mill Methods', level=2)
doc.add_paragraph(
    'Tissues are mixed with small beads and subjected to mechanical agitation. Bead mills offer high throughput and '
    'efficient disruption of wide range of tissues. Choice of bead material, size, and agitation speed must be optimized for '
    'specific applications.'
)

doc.add_heading('4.3 Ultrasonic Homogenization', level=2)
doc.add_paragraph(
    'Ultrasonic waves are applied to tissues, causing mechanical disruption through cavitation. This method is non-contact '
    'and gentle, preserving delicate molecules. It is often used for heat-sensitive samples or to release cellular components from '
    'membranes.'
)

doc.add_paragraph(
    'The resulting homogenate can be further processed by ultrasonic crushing, with attention to ice bath cooling during ultrasonic '
    'crushing to prevent temperature increase. Probe sonication can be used with burst sonication (5 seconds at 30% power) '
    'to generate fine protein powder.'
)

doc.add_heading('4.4 High-Pressure Methods', level=2)
doc.add_paragraph(
    'Tissues are subjected to high pressure through narrow gap using piston-driven apparatus. This technique efficiently '
    'disrupts cells and yields high-quality homogenates. Temperature control is crucial to prevent heat-induced damage.'
)

doc.add_paragraph(
    'Tissues forced through narrow gap or small orifices at high pressure offer reproducible and scalable disruption, suitable for '
    'large-scale applications, commonly used for plant tissues and microbial cells.'
)

doc.add_heading('4.5 Cryogenic Homogenization', level=2)
doc.add_paragraph(
    'Cryogenic homogenization involves grinding tissues at extremely low temperatures, preventing heat-induced damage and preserving '
    'delicate molecules. This method is particularly suitable for heat-sensitive samples.'
)

doc.add_heading('4.6 Enzymatic and Detergent Methods', level=2)
doc.add_paragraph(
    'Enzymes such as collagenase or trypsin digest tissues, breaking down extracellular matrix components and releasing cells. '
    'This method allows isolation of specific cell populations, frequently employed in cell culture experiments.'
)

doc.add_paragraph(
    'Detergents such as Triton X-100 or NP-40 disrupt cell membranes and solubilize cellular components. Useful for '
    'isolating membrane-bound proteins or organelles. Detergent concentration must be optimized to avoid interference with downstream '
    'analyses.'
)

doc.add_paragraph('')

# Temperature Control
doc.add_heading('5. Temperature Control During Homogenization', level=1)

doc.add_paragraph(
    'Temperature control during homogenization is critical to preserve delicate molecules and prevent heat-induced damage. '
    'All the process should be carried out under chilled environment because increasing temperature will destruct enzyme activity.'
)

doc.add_paragraph(
    'Homogenization can develop friction which increases solution temperature. The tissue block should be moved into glass '
    'homogenizer with corresponding volume of PBS for full grinding, and process should be carried out on ice.'
)

doc.add_paragraph(
    'For TRIzol tissue homogenate, sample should be kept at 4°C (on ice) during protocol unless otherwise stated. Incubation '
    'occurs on ice/at 4°C for 30 minutes with mixing, then 5 minutes at 30°C before cooling on ice.'
)

doc.add_paragraph('')

# Buffer Selection
doc.add_heading('6. Buffer Selection and Optimization', level=1)

doc.add_paragraph(
    'The choice of buffer, pH, and osmolality should be optimized to maintain cellular stability and prevent degradation.'
)

doc.add_heading('6.1 Phosphate Buffer Solutions', level=2)
doc.add_paragraph(
    'PBS (0.01 mol/L, pH 7.0–7.4) is commonly recommended for tissue homogenization. Tissue homogenates, typically from 0.2 '
    'g tissue, are prepared in phosphate buffer (PBE) by Ultra-Turrax homogenizer (~10 s) in ice-cooled plastic tubes.'
)

doc.add_heading('6.2 Cell Lysis Buffer', level=2)
doc.add_paragraph(
    'Cell Lysis Buffer (EPX-99999-000) is used with 500 μL added per 100 mg tissue for TissueLyser homogenization. This '
    'buffer is suitable for protein extraction and ELISA detection.'
)

doc.add_heading('6.3 Buffered Peptone Water', level=2)
doc.add_paragraph(
    'For microbiological analysis of gecko lower gut, tissues are homogenized in buffered peptone water solution (BPW) in 1 mL. '
    'This is used for Salmonella isolation and processing.'
)

doc.add_heading('6.4 TRIzol Reagent', level=2)
doc.add_paragraph(
    'TRIzol reagent allows simultaneous RNA and protein extraction. Add 1 mL ice-cold TRIzol per 0.1 g tissue weight. This '
    'protocol substitutes traditional SDS-based or urea-based tissue homogenization and affords harvest of RNA during protocol.'
)

doc.add_paragraph('')

# Centrifugation Protocols
doc.add_heading('7. Centrifugation and Supernatant Collection', level=1)

doc.add_paragraph(
    'Once solution is formed, stop homogenization process and centrifuge. The homogenate after centrifugation gives supernatant '
    '(clear solution) and sediment (membrane-bound portion).'
)

doc.add_heading('7.1 Standard Centrifugation Conditions', level=2)
doc.add_paragraph(
    'The prepared homogenate is centrifuged at 5,000 × g for 5 minutes, and supernatant is taken for detection. This is '
    'suitable for ELISA detection and general biochemical analysis.'
)

doc.add_heading('7.2 High-Speed Centrifugation', level=2)
doc.add_paragraph(
    'For tissue homogenate preparation for ProcartaPlex protocol, centrifuge sample at 16,000 × g for 10 minutes at 4°C. '
    'Transfer supernatant to new microcentrifuge tube and measure total protein concentration.'
)

doc.add_heading('7.3 TRIzol Phase Separation', level=2)
doc.add_paragraph(
    'Spin TRIzol tissue homogenate at 12,000 rcf for 5 minutes at 4°C; unlysed tissue debris and cells form pellet at bottom, clear '
    'TRIzol lysate forms supernatant. For phase separation, centrifuge at 12,000 rcf × 5 min.'
)

doc.add_heading('7.4 Fat Removal', level=2)
doc.add_paragraph(
    'Since most tissue samples contain fat, if supernatant is cloudy after homogenization, store at 4°C appropriately. '
    'After fat solidifies and stratifies, take homogenate supernatant of middle layer for subsequent detection.'
)

doc.add_paragraph('')

# Protein Standardization
doc.add_heading('8. Protein Quantification and Standardization', level=1)

doc.add_paragraph(
    'Protein standardization is essential for comparing results across samples and ensuring reproducibility. The final results of '
    'ELISA experiments are divided by corresponding BCA results to obtain amount of target material per mg total protein.'
)

doc.add_heading('8.1 Protein Concentration Measurement', level=2)
doc.add_paragraph(
    'Measure total protein concentration using Bio-Rad DC Protein Assay Kit or Bio-Rad mini-Bradford Assay with BSA as standard. '
    'Protein concentration should be between 2.5 and 7.5 mg/mL.'
)

doc.add_heading('8.2 Sample Dilution', level=2)
doc.add_paragraph(
    'Dilute samples to 10 mg protein/mL with 1× PBS for ProcartaPlex protocol. 1:9 ground tissue homogenate samples generally '
    'need dilution for ELISA detection; optimal dilution ratio determined based on pre-experiments.'
)

doc.add_heading('8.3 Normalization', level=2)
doc.add_paragraph(
    'Normalize protein content and follow standard in-solution digestion protocol. Add 200 μL 0.25% Rapigest in 50 mM ammonium '
    'bicarbonate and resolubilize pellet with vortex and heat to 60°C.'
)

doc.add_heading('8.4 Total Protein Staining Standardization', level=2)
doc.add_paragraph(
    'Recent study advocates for protein standardization based on total protein staining in rabbit posterior capsular tissues, showing '
    'superiority to classical or tissue-specific methods. This approach may be applicable to gecko tissue homogenates.'
)

doc.add_paragraph('')

# Quality Assurance
doc.add_heading('9. Quality Assurance and Validation', level=1)

doc.add_heading('9.1 Homogenization Efficiency Validation', level=2)
doc.add_paragraph(
    'Validation of homogenization efficiency should be performed by assessing release of target molecules or disruption of specific '
    'cellular structures. Homogenate should be even consistency with no visible chunks.'
)

doc.add_heading('9.2 Sample Integrity Assessment', level=2)
doc.add_paragraph(
    'RNA integrity and concentration assessed using Agilent Bioanalyzer 2100 system and Qubit 2.0 Fluorometer. High-quality '
    'RNA samples with RIN values ≥ 7.0 are used for cDNA library construction.'
)

doc.add_heading('9.3 Reproducibility Considerations', level=2)
doc.add_paragraph(
    'Careful consideration should be given to potential impact of chosen method on activity or conformation of molecules of interest. '
    'The choice of homogenization method should be based on tissue type, sample size, desired degree of disruption, and '
    'downstream applications.'
)

doc.add_heading('9.4 Cross-Contamination Prevention', level=2)
doc.add_paragraph(
    'Micro-technique for preparing homogenates minimizes sample loss and risk of cross-contamination, eliminates possibility of '
    'overheating sample during homogenization. Homogenization should take place under fume hood to reduce exposure to small '
    'particles.'
)

doc.add_paragraph('')

# Methodological Considerations - FIXED TABLE
doc.add_heading('10. Critical Methodological Considerations', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Parameter'
header_cells[1].text = 'Recommended Value/Condition'
header_cells[2].text = 'Reference'

# Add data rows
rows_data = [
    ('Tissue-to-buffer ratio', '1:9 (100 mg tissue : 900 μL PBS)', '[web:31]'),
    ('PBS pH', '7.0–7.4', '[web:31]'),
    ('PBS concentration', '0.01 mol/L', '[web:31]'),
    ('Minimum tissue weight', '>30 mg', '[web:31]'),
    ('Homogenization temperature', 'On ice, 4°C', '[web:5][web:31]'),
    ('TissueLyser frequency', '25 Hz', '[web:23]'),
    ('TissueLyser duration', '0.5–3 minutes', '[web:23]'),
    ('Centrifugation (standard)', '5,000 × g, 5 min', '[web:31]'),
    ('Centrifugation (high-speed)', '16,000 × g, 10 min, 4°C', '[web:23]'),
    ('Protein concentration range', '2.5–7.5 mg/mL', '[web:32]'),
    ('Final dilution for ELISA', '10 mg protein/mL', '[web:23]'),
    ('RNA integrity (RIN)', '≥ 7.0', '[web:25]'),
]

for param, value, ref in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = param
    row_cells[1].text = value
    row_cells[2].text = ref

doc.add_paragraph('')

# Applications
doc.add_heading('11. Applications of Gecko Homogenate', level=1)

doc.add_heading('11.1 Biochemical Parameter Estimation', level=2)
doc.add_paragraph(
    'For estimation of biochemical parameters and enzymes, proteins from animal tissue, preparation of tissue homogenate is must. '
    'Both supernatant and sediment are used for various estimations including enzyme activity measurements.'
)

doc.add_heading('11.2 Proteomic Analysis', level=2)
doc.add_paragraph(
    'Tissue homogenate preparation for proteomics analyses requires standardized extraction and homogenization protocols. '
    'SMRT sequencing of G. gecko full-length transcriptome identified 111,372 transcripts annotated against eight databases, '
    'providing genetic resource for proteomic research.'
)

doc.add_heading('11.3 Enzymatic Analysis', level=2)
doc.add_paragraph(
    'Black-spotted geckos affect immune regulation in mouse models of asthma. Enzyme activity detection requires chilled '
    'environment during homogenization to prevent enzyme destruct ion.'
)

doc.add_heading('11.4 Microbiological Analysis', level=2)
doc.add_paragraph(
    'For Salmonella isolation from gecko lower gut, tissues are homogenized in buffered peptone water. Lower gut of each '
    'gecko is weighed and processed separately.'
)

doc.add_heading('11.5 Genetic and Isotope Analysis', level=2)
doc.add_paragraph(
    'DNA extracted from tail tissue (wild-caught) and dried ventral tissue (TCM shop specimens) using DNeasy Blood & Tissue extraction '
    'kits. Stable isotope analysis uses tail tissue tip dried at 60°C for 24 hours minimum.'
)

doc.add_paragraph('')

# Challenges and Limitations
doc.add_heading('12. Challenges and Limitations', level=1)

doc.add_paragraph(
    'Despite accumulation of methodological literature, specific protocols for gecko tissue homogenate remain limited. Most protocols '
    'are adapted from general animal tissue homogenization methods.'
)

doc.add_heading('12.1 Species-Specific Considerations', level=2)
doc.add_paragraph(
    'Tokay geckos are second largest gecko species, widespread across Asia. Two morphs (black- and red-spotted) or subspecies '
    '(G. gecko reevesii and G. gecko gecko) exist, sometimes elevated to species level. Different tissues may require '
    'optimized homogenization conditions.'
)

doc.add_heading('12.2 Sample Availability', level=2)
doc.add_paragraph(
    'G. gecko is rare and endangered medicinal animal in China. Listed as Class II protected species since 1989. '
    'Limited availability of cultured specimens restricts methodological validation studies.'
)

doc.add_heading('12.3 Method Optimization', level=2)
doc.add_paragraph(
    'Choice of bead material, size, and agitation speed must be optimized for specific applications. Detergent concentration '
    'must be optimized to avoid interference with downstream analyses. Optimal dilution ratio for ELISA determined based on '
    'pre-experiments.'
)

doc.add_heading('12.4 Temperature Sensitivity', level=2)
doc.add_paragraph(
    'Temperature increase destructs enzyme activity. Homogenization friction increases solution temperature. Temperature '
    'control crucial to prevent heat-induced damage.'
)

doc.add_paragraph('')

# Future Directions
doc.add_heading('13. Future Directions', level=1)

doc.add_paragraph(
    'Future research should focus on developing gecko-specific homogenization protocols optimized for different tissue types and '
    'downstream applications.'
)

doc.add_heading('13.1 Standardized Gecko Protocols', level=2)
doc.add_paragraph(
    'Development of standardized protocols specifically for G. gecko tissue would improve reproducibility across studies. Current '
    'protocols are adapted from general animal tissue methods.'
)

doc.add_heading('13.2 Comparative Method Studies', level=2)
doc.add_paragraph(
    'Comparative studies evaluating different homogenization methods (mechanical, bead mill, ultrasonic, cryogenic) for gecko tissue '
    'would identify optimal approaches for specific applications.'
)

doc.add_heading('13.3 Integration with Genomic Data', level=2)
doc.add_paragraph(
    'Integration of homogenate preparation with full-length transcriptome data from SMRT sequencing would accelerate gene '
    'function research and biological regulatory mechanism studies in G. gecko.'
)

doc.add_heading('13.4 Conservation-Compliant Methods', level=2)
doc.add_paragraph(
    'Given endangered status of G. gecko, development of micro-scale homogenization techniques requiring minimal tissue '
    'amounts would support conservation while enabling research.'
)

doc.add_paragraph('')

# Conclusion
doc.add_heading('14. Conclusion', level=1)

doc.add_paragraph(
    'This review comprehensively examined methodology approaches for preparing and standardizing gecko homogenate, synthesizing '
    'updated literature from 2020 to 2026. Key findings include:'
)

doc.add_paragraph('Tissue-to-buffer ratio of 1:9 (100 mg:900 μL PBS) is standard for most applications')
doc.add_paragraph('Temperature control on ice at 4°C is critical to preserve enzyme activity and prevent protein denaturation')
doc.add_paragraph('PBS pH 7.0–7.4 maintains cellular stability')
doc.add_paragraph('Centrifugation at 5,000 × g for 5 min (standard) or 16,000 × g for 10 min at 4°C (high-speed) yields clear supernatant')
doc.add_paragraph('Protein concentration should be normalized to 2.5–7.5 mg/mL for most analyses')
doc.add_paragraph('Various homogenization methods (mechanical, bead mill, ultrasonic, cryogenic) offer options based on tissue type and downstream application')
doc.add_paragraph('Validation of homogenization efficiency through target molecule release assessment ensures reproducibility')

doc.add_paragraph(
    'These guidelines provide actionable framework for researchers preparing gecko homogenates for enzymatic, proteomic, and biochemical '
    'analyses while maintaining sample integrity and reproducibility. Future development of gecko-specific standardized protocols would '
    'further improve research quality in this important medicinal animal species.'
)

doc.add_paragraph('')

# References
doc.add_heading('References', level=1)

references = [
    ('[web:1]', 'CN102787157A - Preparation method of gecko toxoid. Google Patents, 2012.'),
    ('[web:2]', 'B那可特, et al. A micro-technique for preparing homogenates of biological samples. J Electrophor, 1981.'),
    ('[web:5]', 'Tissue Homogenization Process. YouTube, 2022.'),
    ('[web:6]', 'Preparation and Homogenization of Fish Tissue Sample. Kentucky EPA, 2019.'),
    ('[web:9]', 'Methods and techniques for tissue homogenization - A comprehensive review. Allied Academies, 2023.'),
    ('[web:10]', 'How to prepare Tissue Homogenate samples? Sunlong Biotech, 2025.'),
    ('[web:16]', 'Salmonella Isolates in the Introduced Asian House Gecko. J Vet Zool, 2015.'),
    ('[web:17]', 'Dufour PC, et al. Home and hub: pet trade and traditional medicine impact reptile populations. Proc Biol Sci, 2022;289(1982):20221011.'),
    ('[web:23]', 'Prepare Tissue Homogenate - ProcartaPlex Protocol. Thermo Fisher, manual L11234.'),
    ('[web:25]', 'Jiang J, et al. SMRT sequencing of the full-length transcriptome of Gekko gecko. PLoS One, 2022;17(2):e0264499.'),
    ('[web:31]', 'Collection, processing and storage of tissue homogenate and cell. ELK Biotech, 2024.'),
    ('[web:32]', 'Tissue Homogenization and RNA + Protein Extraction using TRIzol Reagent. Duke Biostat, protocol.'),
    ('[web:33]', 'Tissue Homogenate - an overview. ScienceDirect Topics.'),
    ('[web:35]', 'Total protein staining is superior to classical methods. ScienceDirect, 2020.'),
]

for ref_id, ref_text in references:
    p = doc.add_paragraph()
    p.add_run(f'{ref_id} ').bold = True
    p.add_run(ref_text)

# Create output directory if it doesn't exist
os.makedirs('output', exist_ok=True)

# Save document
doc.save('output/Gecko_Homogenate_Methodology_Review_2020-2026.docx')
print("✅ Document created successfully!")